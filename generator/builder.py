import os
from datetime import datetime
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import LETTER
from jinja2 import Template

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "generator", "templates", "base_fir.txt")

def _load_template() -> Template:
    with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
        return Template(f.read())

def generate_fir_text(inputs: dict) -> str:
    """
    Renders the FIR text using a Jinja template.
    All fields are treated as user-supplied content; no legal advice is given.
    """
    # basic normalization
    inputs = {k: (v.strip() if isinstance(v, str) else v) for k, v in inputs.items()}
    inputs.setdefault("generated_on", datetime.now().strftime("%Y-%m-%d %H:%M"))
    template = _load_template()
    return template.render(**inputs).strip()

def build_docx(text: str, filepath: str):
    doc = Document()
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")  # blank line
        else:
            doc.add_paragraph(line)
    doc.save(filepath)

def build_pdf(text: str, filepath: str):
    styles = getSampleStyleSheet()
    story = []
    for line in text.splitlines():
        if line.strip() == "":
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(line, styles["Normal"]))
    doc = SimpleDocTemplate(filepath, pagesize=LETTER, leftMargin=50, rightMargin=50, topMargin=50, bottomMargin=50)
    doc.build(story)
